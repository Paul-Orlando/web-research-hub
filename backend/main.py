import base64
import io
import json
import os
import re
from urllib.parse import urlparse

import httpx
from dotenv import load_dotenv

load_dotenv()

MCP_SERVER_URL = os.getenv("MCP_SERVER_URL")

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from sse_starlette.sse import EventSourceResponse

from agents.orchestrator import run_subtask
from agents.planner import run_planner
from agents.summarizer import run_summarizer
from models.schemas import DepthMode, ResearchRequest
from session_store import session_store

app = FastAPI(title="Web Research Agent")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── helpers ──────────────────────────────────────────────────────────────────

def _links_to_domain(text: str) -> str:
    """Replace [label](url) with label (domain) for plain-text export."""
    def replace(m: re.Match) -> str:
        label, url = m.group(1), m.group(2)
        try:
            domain = urlparse(url).netloc.removeprefix("www.")
        except Exception:
            domain = ""
        return f"{label} ({domain})" if domain else label
    return re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", replace, text)


def _sse(event: str, data: dict) -> dict:
    return {"event": event, "data": json.dumps(data)}


def _generate_pdf(markdown_text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    def safe(text: str) -> str:
        return text.encode("ascii", "ignore").decode("ascii")

    def write(size: int, bold: bool, text: str, line_height: float = 7.0) -> None:
        pdf.set_font("Helvetica", style="B" if bold else "", size=size)
        pdf.multi_cell(0, line_height, safe(text), new_x="LMARGIN", new_y="NEXT")

    for line in markdown_text.split("\n"):
        if line.startswith("# "):
            pdf.ln(2)
            write(18, True, line[2:], 10)
            pdf.ln(1)
        elif line.startswith("## "):
            pdf.ln(2)
            write(14, True, line[3:], 9)
            pdf.ln(1)
        elif line.startswith("### "):
            pdf.ln(1)
            write(12, True, line[4:], 8)
        elif line.startswith(("- ", "* ")):
            text = _links_to_domain(line[2:])
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            write(11, False, "- " + text)
        elif line.strip():
            text = _links_to_domain(line)
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            write(11, False, text)
        else:
            pdf.ln(3)

    return bytes(pdf.output())


def _generate_docx(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in markdown_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ", "* ")):
            text = _links_to_domain(line[2:])
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
        elif line.strip():
            text = _links_to_domain(line)
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


async def _mcp_export_report(content: str, fmt: str, title: str) -> bytes:
    payload = {
        "jsonrpc": "2.0",
        "method": "tools/call",
        "id": 1,
        "params": {
            "name": "export_report",
            "arguments": {
                "content": content,
                "format": fmt,
                "title": title,
            },
        },
    }
    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            f"{MCP_SERVER_URL}/mcp",
            json=payload,
            headers={"Content-Type": "application/json"},
        )
        response.raise_for_status()
        data = response.json()

    text = data["result"]["content"][0]["text"]
    parsed = json.loads(text)
    return base64.b64decode(parsed["content_base64"])


# ── endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/research")
async def research(request: ResearchRequest):
    async def event_stream():
        session = session_store.get_or_create(request.session_id)
        history = session["history"]
        all_results = []

        try:
            # Stage 1: Planning
            yield _sse("stage", {"stage": "planning", "message": "Planning search strategy..."})

            plan = await run_planner(request.query, request.depth_mode, history)
            subtasks = plan.subtasks

            yield _sse("subtasks", {
                "count": len(subtasks),
                "subtasks": [s.model_dump() for s in subtasks],
            })

            # Stage 2: Searching
            yield _sse("stage", {
                "stage": "searching",
                "message": f"Searching {len(subtasks)} sources...",
            })

            for idx, subtask in enumerate(subtasks):
                results = await run_subtask(subtask, request.depth_mode)
                all_results.extend(results)

                sources = [
                    {"url": r.source, "publish_date": r.publish_date, "title": r.title}
                    for r in results
                    if r.source
                ]
                yield _sse("subtask_complete", {
                    "id": subtask.id,
                    "query": subtask.query,
                    "sources": sources,
                    "index": idx,
                    "total": len(subtasks),
                })

            # Stage 3: Summarizing
            yield _sse("stage", {"stage": "synthesizing", "message": "Synthesizing report..."})

            report_markdown = await run_summarizer(all_results, request.query, history)

            # Persist to session
            session_sources = [
                {"url": r.source, "publish_date": r.publish_date, "title": r.title}
                for r in all_results
                if r.source
            ]
            session_store.update(
                request.session_id,
                request.query,
                report_markdown,
                session_sources,
            )

            yield _sse("report", {"markdown": report_markdown})
            yield _sse("done", {})

        except Exception as exc:
            yield _sse("error", {"message": str(exc)})

    return EventSourceResponse(event_stream())


@app.get("/session/{session_id}/sources")
async def get_sources(session_id: str):
    return session_store.get_sources(session_id)


@app.get("/session/{session_id}/history")
async def get_history(session_id: str):
    return session_store.get_history(session_id)


@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    session_store.delete(session_id)
    return {"deleted": session_id}


@app.get("/download-report/{session_id}/{fmt}")
async def download_report(session_id: str, fmt: str):
    history = session_store.get_history(session_id)
    if not history:
        raise HTTPException(status_code=404, detail="No report found for this session")

    report = history[-1]["report"]
    title = history[-1].get("query", "Research Report")

    if fmt == "md":
        return Response(
            content=report,
            media_type="text/markdown",
            headers={"Content-Disposition": "attachment; filename=report.md"},
        )
    elif fmt == "pdf":
        try:
            content = await _mcp_export_report(report, "pdf", title)
        except Exception:
            content = _generate_pdf(report)
        return Response(
            content=content,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=report.pdf"},
        )
    elif fmt == "docx":
        try:
            content = await _mcp_export_report(report, "docx", title)
        except Exception:
            content = _generate_docx(report)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=report.docx"},
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unknown format: {fmt}")
